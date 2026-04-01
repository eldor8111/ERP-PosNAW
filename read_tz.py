import sys
import subprocess
try:
    import docx  # type: ignore
except ImportError:
    subprocess.check_call([sys.executable, "-m", "pip", "install", "python-docx"])
    import docx  # type: ignore

if hasattr(sys.stdout, 'reconfigure'):
    sys.stdout.reconfigure(encoding='utf-8')  # type: ignore

doc_path = r"C:\Users\ELDORBEK ABDUALIMOV\Desktop\TZPOSERP.docx"
doc = docx.Document(doc_path)
for para in doc.paragraphs:
    if para.text.strip():
        print(para.text)
for table in doc.tables:
    print("--- TABLE ---")
    for row in table.rows:
        print(" | ".join([cell.text.replace("\n", " ").strip() for cell in row.cells]))
